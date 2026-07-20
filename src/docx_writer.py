from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .docx_layout import add_title_block, add_watermark, is_kalam_installed, setup_footer, setup_header
from .dtp_parser import parse_dtp_note
from .image_regions import crop_region, detect_regions
from .image_tools import extract_pw_logo, smart_crop_image
from .math_renderer import (
    ARROW as _ARROW,
    BAR as _BAR,
    COMMON_FORMULAS as _COMMON_FORMULAS,
    HAT as _HAT,
    LATEX_SYMBOLS as _LATEX_SYMBOLS,
    MATH_FONT,
    MATH_TAG_RE,
    PRE_HAT as _PRE_HAT,
    RENDER_OMML,
    accent as _accent,
    add_block_math,
    add_inline_math,
    has_math_tag,
    split_math_segments,
)
from .models import SlideData

BODY_FONT = "Kalam"
BODY_SIZE = Pt(16)
HEADING_SIZE = Pt(18)
DARK_GREY = RGBColor(0x00, 0x00, 0x99)   # body text (#000099)
BOLD_COLOR = RGBColor(0xE9, 0x71, 0x32)  # bold words (#E97132)
BULLET_PREFIXES = ("•", "·", "-", "–")
BULLET_MARKERS = ("•", "◦", "▪", "–")
CHAR_SPACING_TWIPS = 30  # Expanded character spacing = 1.5 pt (1 pt = 20 twips)


def _set_char_spacing(run, twips: int = CHAR_SPACING_TWIPS):
    rpr = run._element.get_or_add_rPr()
    sp = rpr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        rpr.append(sp)
    sp.set(qn("w:val"), str(twips))


def _set_no_same_style_spacing(p):
    ppr = p._element.get_or_add_pPr()
    if ppr.find(qn("w:contextualSpacing")) is None:
        ppr.append(OxmlElement("w:contextualSpacing"))


def _set_run_font(run, *, size=BODY_SIZE, bold=False, italic=False, underline=False, color=DARK_GREY, font=BODY_FONT):
    # `font` is overridden to a maths font for formula runs: Kalam has no
    # combining arrow/hat glyphs, so forcing it is what erases vectors and hats.
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = color
    _set_char_spacing(run)


def _set_paragraph_base(p, *, bullet: bool = False, indent_level: int = 0):
    try:
        p.style = "Body Text"
    except Exception:
        pass
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_no_same_style_spacing(p)
    if bullet:
        # Each nesting level shifts the bullet right by ~0.7 cm while keeping the
        # hanging indent so wrapped lines align under the text, not the marker.
        base = 0.6 + 0.7 * max(0, indent_level)
        p.paragraph_format.left_indent = Cm(base)
        p.paragraph_format.first_line_indent = Cm(-0.6)
    else:
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)


# Order matters: try ***bold-italic*** before **bold** before *italic* so the
# alternation consumes the longest marker first at each position.
_MD_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")

_GREEK = {
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ", "epsilon": "ε",
    "zeta": "ζ", "eta": "η", "theta": "θ", "iota": "ι", "kappa": "κ",
    "lambda": "λ", "mu": "μ", "nu": "ν", "xi": "ξ", "rho": "ρ",
    "sigma": "σ", "tau": "τ", "upsilon": "υ", "phi": "φ", "chi": "χ",
    "psi": "ψ", "omega": "ω", "pi": "π",
}
_GREEK_UPPER = {
    "delta": "Δ", "omega": "Ω", "sigma": "Σ", "phi": "Φ", "theta": "Θ",
    "lambda": "Λ", "pi": "Π", "gamma": "Γ", "psi": "Ψ", "xi": "Ξ",
}
_GREEK_RE = re.compile(r"\b(" + "|".join(_GREEK) + r")\b", re.I)

# Chars that mark a maths context. A bare greek WORD (e.g. "theta") is only
# turned into a symbol when it touches one of these — so science prose like
# "alpha particle", "beta decay", "gamma rays", "sigma bond" or "pi bond"
# keeps its English words instead of becoming α/β/γ/σ/π.
_MATH_NEIGHBOUR = set("=+-*/^_()[]{}×·√±≤≥<>∫∑∏°|⇒→∝≈≡0123456789"
                      "∈∉∋⊂⊆⊃⊇∪∩∖∅∀∃∄⇐⇔↔↦∴∵∣⟨⟩∮∂∇⊥∠∥←")


def _greek_sub(m):
    """Bare greek word -> symbol, but ONLY in a maths context (see above).
    'theta' -> θ, 'Delta' -> Δ (capitalised word => capital letter)."""
    s, i, j = m.string, m.start(), m.end()
    before = s[:i].rstrip()[-1:]
    after = s[j:].lstrip()[:1]
    if before not in _MATH_NEIGHBOUR and after not in _MATH_NEIGHBOUR:
        return m.group(1)                       # prose word -> leave untouched
    word = m.group(1)
    low = word.lower()
    if word[0].isupper() and low in _GREEK_UPPER:
        return _GREEK_UPPER[low]
    return _GREEK[low]


# The LaTeX symbol table, combining accents (_HAT/_ARROW/_BAR) and the chemical
# formula whitelist all live in math_renderer now, so the tagged-maths path and
# this plain-text normaliser can never drift apart. They are imported above.


# Always-safe display math: $$..$$, \(..\), \[..\]. Inline $..$ is handled
# separately so currency ("$5 and $10") is never mistaken for maths.
_DISPLAY_MATH_RE = re.compile(r"\$\$(.+?)\$\$|\\\((.+?)\\\)|\\\[(.+?)\\\]", re.S)
_INLINE_DOLLAR_RE = re.compile(r"\$([^$\n]{1,300}?)\$")
# \frac{a}{b}, \sqrt{x}, \sqrt[n]{x}, \vec{v}, \hat{i}, \text{x}
_FRAC_RE = re.compile(r"\\[dt]?frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}")
_SQRT_BRACE_RE = re.compile(r"\\sqrt\s*(?:\[([^\]]*)\])?\s*\{([^{}]*)\}")
_ACCENT_RE = re.compile(r"\\(vec|overrightarrow|hat|widehat|bar|overline)\s*(?:\{([^{}]*)\}|([A-Za-z0-9]))")
_WRAP_RE = re.compile(r"\\(?:text|textbf|textit|mathrm|mathbf|mathit|mathsf|operatorname)\s*\{([^{}]*)\}")
_LATEX_CMD_RE = re.compile(r"\\([A-Za-z]+)")

# A "*" that means multiplication: either tight between operands ("2*v_A") or
# spaced on BOTH sides ("a * b"). Markdown emphasis never looks like that — the
# opening "*" of *italic* has no operand before it and no space after it — so
# **bold** / *italic* survive for _MD_RE. This MUST run before the markdown
# split, otherwise "2*v_A*v_B" is read as italics and the asterisks are eaten.
_MULT_RE = re.compile(r"(?<=[0-9A-Za-z_)\]])(?:\*|\s+\*\s+)(?=[0-9A-Za-z_(\[√])")

# Unit-vector / hat notation written as "i^", "k^" (a letter then a caret that
# is NOT the start of a superscript). Turned into î / k̂ etc. The (?![+-]\s*\d)
# guard keeps a NEGATIVE exponent like "s^-2" out of the hat rule.
_HAT_CARET_RE = re.compile(r"([A-Za-z])\^(?![A-Za-z0-9{]|[+-]\s*\d)")

# base + sub/superscript token. The token is a SINGLE character class so
# "H_2O" subscripts only the "2" (not "2O"), while "v_AB" still subscripts "AB".
# An optional leading sign captures negative/positive exponents: 10^-3, x^+2.
_SCRIPT_RE = re.compile(r"([_^])(\{[^}]{1,40}\}|[+-]?\d+|[A-Za-z]+)")

# Common chemical formulas — subscripted safely (whitelist => zero false
# positives on things like "A4", "MP3", "Class 11"). Insert "_" before each
# digit run that follows a letter, e.g. "CO2" -> "CO_2", handled downstream.
_COMMON_FORMULA_RE = re.compile(
    r"(?<![A-Za-z0-9])(" + "|".join(re.escape(f) for f in sorted(_COMMON_FORMULAS, key=len, reverse=True))
    + r")(?![A-Za-z0-9])"
)


def _accent_sub(m):
    kind, braced, single = m.group(1), m.group(2), m.group(3)
    base = braced if braced is not None else (single or "")
    comb = {"hat": _HAT, "widehat": _HAT, "vec": _ARROW, "overrightarrow": _ARROW,
            "bar": _BAR, "overline": _BAR}[kind]
    return _accent(base, comb)


def _unwrap_display(m):
    return next(g for g in m.groups() if g is not None)


def _unwrap_inline_dollar(m):
    """Only treat $..$ as maths when the content actually looks like maths,
    so prose amounts like "$5 and $10" keep their dollar signs."""
    inner = m.group(1)
    if re.search(r"[\\_^=]", inner) or re.search(r"[A-Za-z]\s*[+\-/×·]\s*[A-Za-z0-9]", inner):
        return inner
    return m.group(0)


def _normalize_math(text: str) -> str:
    """Normalise any maths notation the model emits into plain readable symbols.

    Handles LaTeX (\\frac, \\sqrt, \\theta, \\hat, $...$), ASCII (sqrt(), *, i^,
    theta, <=) and already-correct Unicode, so the result only ever uses real
    symbols plus '_'/'^' markers that _add_scripted_runs turns into
    sub/superscripts. Runs BEFORE the markdown split so formula asterisks are
    never read as italics.
    """
    text = text.replace("\\\\", " ")
    text = re.sub(r"\\([%&#{}$_])", r"\1", text)      # unescape \% \& \{ \$ ...
    text = re.sub(r"`([^`]*)`", r"\1", text)          # `v_AB` code ticks
    text = _DISPLAY_MATH_RE.sub(_unwrap_display, text)
    text = _INLINE_DOLLAR_RE.sub(_unwrap_inline_dollar, text)
    # Strip ONLY the bare \left / \right delimiter commands — the negative
    # lookahead keeps \rightarrow, \leftarrow, \leftrightarrow intact for the
    # symbol table below (without it, "\rightarrow" became the word "arrow").
    text = re.sub(r"\\left(?![A-Za-z])\s*|\\right(?![A-Za-z])\s*", "", text)
    text = re.sub(r"\\[,;:!> ]", " ", text)
    # Structural LaTeX, innermost-first (a few passes handles simple nesting).
    for _ in range(4):
        new = _WRAP_RE.sub(r"\1", text)
        new = _ACCENT_RE.sub(_accent_sub, new)
        new = _FRAC_RE.sub(r"(\1)/(\2)", new)
        new = _SQRT_BRACE_RE.sub(lambda m: f"√({m.group(2)})", new)
        if new == text:
            break
        text = new
    # Remaining LaTeX commands -> symbols. An unknown command just loses its
    # backslash (\cos -> cos), since a stray backslash must never reach the page.
    text = _LATEX_CMD_RE.sub(lambda m: _LATEX_SYMBOLS.get(m.group(1), m.group(1)), text)
    # ASCII maths.
    text = re.sub(r"\bsqrt\s*\(", "√(", text, flags=re.I)
    text = _MULT_RE.sub("×", text)
    # Mark subscripts in known bare chemical formulas ("CO2" -> "CO_2") so they
    # render even when the model forgot the "_" the prompt asked for.
    text = _COMMON_FORMULA_RE.sub(
        lambda m: re.sub(r"(?<=[A-Za-z)])(\d+)", r"_\1", m.group(1)), text)
    text = _HAT_CARET_RE.sub(lambda m: _accent(m.group(1), _HAT), text)
    text = re.sub(r"\^\s*\{?\s*°\s*\}?", "°", text)   # 30^\circ / 30^{°} -> 30°
    text = _GREEK_RE.sub(_greek_sub, text)
    for src, dst in (("<=", "≤"), (">=", "≥"), ("!=", "≠"), ("+/-", "±"),
                     ("->", "→"), ("&gt;", ">"), ("&lt;", "<"), ("&amp;", "&")):
        text = text.replace(src, dst)
    return text


def _add_scripted_runs(p, clean: str, *, size, color, bold, italic, underline, highlight, font=BODY_FONT):
    """Add `clean` as runs, rendering H_2O / v_AB / x^2 as real sub/superscripts.

    A letter subscript/superscript only fires when its base is a lone atom (a
    single symbol not sitting inside a word), so ordinary prose that happens to
    contain an underscore (e.g. "slides_raw.json") is left alone. Digit and
    braced tokens always apply (chemistry counts, exponents, x_{net})."""

    def _emit(chunk: str, *, sub: bool = False, sup: bool = False):
        if not chunk:
            return
        run = p.add_run(chunk)
        _set_run_font(run, size=size, bold=bold, italic=italic, underline=underline, color=color, font=font)
        if sub:
            run.font.subscript = True
        if sup:
            run.font.superscript = True
        if highlight is not None:
            run.font.highlight_color = highlight

    pos = 0
    for m in _SCRIPT_RE.finditer(clean):
        token = m.group(2)
        is_letters = token[0].isalpha()
        prev = clean[m.start() - 1] if m.start() > 0 else ""
        prev2 = clean[m.start() - 2] if m.start() >= 2 else ""
        # A letter subscript on a base that is itself inside a word is almost
        # certainly prose (file_name, slides_raw) -> emit literally.
        if is_letters and prev.isalnum() and prev2.isalnum():
            continue
        _emit(clean[pos:m.start()])
        if token.startswith("{"):
            token = token[1:-1]
        _emit(token, sub=m.group(1) == "_", sup=m.group(1) == "^")
        pos = m.end()
    _emit(clean[pos:])


def _math_fallback_writer(paragraph, text: str, *, size=None, color=None):
    """Write a formula's Unicode fallback using the maths font.

    Any ``^``/``_`` marker latex_to_unicode could not map to a Unicode glyph is
    turned into a REAL Word superscript/subscript run here, so nothing is lost.
    """
    _add_scripted_runs(
        paragraph, text, size=size or BODY_SIZE, color=color if color is not None else DARK_GREY,
        bold=False, italic=False, underline=False, highlight=None, font=MATH_FONT,
    )


def _add_markdown_runs(p, text: str, *, size=BODY_SIZE, base_color=DARK_GREY, default_bold: bool = False, underline: bool = False, highlight=None, math_mode: str = RENDER_OMML, math_stats: dict | None = None):
    """Render one line: protected maths tags become equations, the rest is prose."""
    for kind, payload in split_math_segments(text):
        if kind != "text":
            used = add_inline_math(
                p, payload, mode=math_mode, size=size, color=base_color,
                fallback_writer=_math_fallback_writer,
            )
            if math_stats is not None:
                math_stats[used] = math_stats.get(used, 0) + 1
            continue
        for part in _MD_RE.split(_normalize_math(payload)):
            if part == "":
                continue
            bold = default_bold
            italic = False
            clean = part
            if part.startswith("***") and part.endswith("***") and len(part) >= 6:
                bold = italic = True
                clean = part[3:-3]
            elif part.startswith("**") and part.endswith("**") and len(part) >= 4:
                bold = True
                clean = part[2:-2]
            elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
                italic = True
                clean = part[1:-1]
            _add_scripted_runs(
                p, clean, size=size,
                color=BOLD_COLOR if bold and not default_bold else base_color,
                bold=bold, italic=italic, underline=underline, highlight=highlight,
            )


def _is_dtp(line: str) -> bool:
    return "note to dtp" in line.lower()


def _is_bullet(line: str) -> bool:
    s = line.strip()
    if s.startswith(BULLET_PREFIXES):
        return True
    # Markdown-style "* item" bullet (a single asterisk + space), NOT "**bold**".
    if re.match(r"^\*\s+\S", s):
        return True
    return bool(re.match(r"^\s*\d+[.)]\s+", line))


# Covers both ASCII/LaTeX notation AND the Unicode symbols this writer itself
# produces (set/logic/arrows), so a standalone equation line like "x ∈ A ∪ B"
# or "A → B" is never styled as a green heading.
_FORMULA_HINT_RE = re.compile(
    r"[=√×÷≤≥≠±∝∫∑∈∉∋⊂⊆⊃⊇∪∩∖∅∀∃∄→←⇒⇐⇔↔↦∴∵∮∂∇⊥∠∥⟨⟩]"
    r"|[A-Za-z0-9]\^|[A-Za-z0-9]_[A-Za-z0-9{]|\\[A-Za-z]+|\$"
)


def _looks_like_formula(s: str) -> bool:
    """A standalone equation line (e.g. 'v_AB = v_A - v_B' or 'F = m*a') must not
    be rendered as a green highlighted heading. A protected maths tag always
    counts, so a tagged equation on its own line never becomes a heading."""
    return has_math_tag(s) or bool(_FORMULA_HINT_RE.search(s))


def _is_heading(line: str) -> bool:
    s = line.strip()
    if not s or _is_bullet(s) or _is_dtp(s):
        return False
    if len(s) > 140:
        return False
    if s.startswith("#"):
        return True
    if s.endswith(".") or s.endswith("।"):
        return False
    if _looks_like_formula(s):
        return False
    return True


def _add_heading(doc: Document, line: str, *, math_mode: str = RENDER_OMML, math_stats: dict | None = None):
    p = doc.add_paragraph()
    _set_paragraph_base(p)
    clean = line.strip().strip("#").strip()
    _add_markdown_runs(
        p,
        clean,
        size=HEADING_SIZE,
        base_color=RGBColor(0, 0, 0),
        default_bold=True,
        underline=True,
        highlight=WD_COLOR_INDEX.BRIGHT_GREEN,
        math_mode=math_mode,
        math_stats=math_stats,
    )
    return p


def _add_blank_line(doc: Document):
    p = doc.add_paragraph()
    _set_paragraph_base(p)
    return p


def _add_body(doc: Document, line: str, *, bullet: bool = False, indent_level: int = 0, math_mode: str = RENDER_OMML, math_stats: dict | None = None):
    p = doc.add_paragraph()
    _set_paragraph_base(p, bullet=bullet, indent_level=indent_level)
    if bullet:
        s = line.strip()
        # Normalise bullet markers and make nested bullets visibly different.
        m = re.match(r"^(\*|•|·|-|–)\s+(.*)$", s)
        if m:
            marker = BULLET_MARKERS[min(max(0, indent_level), len(BULLET_MARKERS) - 1)]
            s = f"{marker}\t{m.group(2)}"
        _add_markdown_runs(p, s, math_mode=math_mode, math_stats=math_stats)
    else:
        _add_markdown_runs(p, line.strip(), math_mode=math_mode, math_stats=math_stats)
    return p


def _add_dtp(doc: Document, line: str, policy: str):
    if policy == "hide_note_insert_image":
        return None
    p = doc.add_paragraph()
    _set_paragraph_base(p)
    run = p.add_run(line.strip())
    _set_run_font(run, size=BODY_SIZE, bold=False, color=RGBColor(255, 0, 0))
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def _has_alpha(image_path: Path) -> bool:
    try:
        from PIL import Image

        with Image.open(image_path) as im:
            return im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info)
    except Exception:
        return False


def _pick_dtp_image(slide, run_dir: Path, region_cache: dict, used: dict):
    """Choose which image this DTP note should get.

    A slide often holds several diagrams and several notes point at the same
    slide. Give each note its own region (in reading order) so different parts
    land at their own place, instead of re-pasting the whole slide every time.
    Returns (path, already_cropped) or (None, False) when there is nothing left
    to add — the caller then skips rather than repeating an image.
    """
    src = Path(slide.image_path)
    if src not in region_cache:
        region_cache[src] = detect_regions(src)
    regions = region_cache[src]
    idx = used.get(src, 0)

    if len(regions) > 1:
        if idx >= len(regions):
            return None, False           # every region used -> don't repeat one
        crop = crop_region(src, regions[idx], run_dir / "inserted_images",
                           f"r{idx + 1}")
        used[src] = idx + 1
        if crop:
            return crop, True
        return (src, False) if idx == 0 else (None, False)

    # One indivisible diagram: insert it once, never again.
    if idx:
        return None, False
    used[src] = 1
    return src, False


def _insert_slide_image(doc: Document, image_path: Path, run_dir: Path, mode: str = "smart_crop", already_cropped: bool = False):
    if not image_path or not Path(image_path).exists():
        return False
    img = Path(image_path)
    # Skip smart-crop for transparent (AI-redrawn) images: cropping flattens the
    # alpha channel and we'd lose the see-through background over the watermark.
    # Region crops are already tight, so they are inserted as-is.
    if mode == "smart_crop" and not already_cropped and not _has_alpha(img):
        img = smart_crop_image(img, run_dir / "inserted_images")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(str(img), width=Inches(5.8))
        return True
    except Exception:
        return False


# A line that is nothing but a display equation.
_BLOCK_ONLY_RE = re.compile(r"^\s*\[\[MATH_BLOCK:\s*(.*?)\]\]\s*$", re.S)


def _normalize_math_tags(text: str) -> str:
    """Put every [[MATH_BLOCK]] on its own line and flatten newlines inside tags.

    The model writes block maths across several lines; collapsing it first means
    the line-by-line writer below sees one complete equation per line.
    """
    def repl(m):
        latex = " ".join(m.group(2).split())
        if m.group(1) == "BLOCK":
            return f"\n[[MATH_BLOCK: {latex}]]\n"
        return f"[[MATH_INLINE: {latex}]]"
    return MATH_TAG_RE.sub(repl, text)


def write_notes_docx(notes_text: str, output_path: Path, slides: list[SlideData], *, run_dir: Path, image_insert_mode: str = "smart_crop", dtp_note_policy: str = "keep_note_and_insert_image", subject: str = "", chapter_title: str = "", math_render_mode: str = RENDER_OMML) -> tuple[Path, list[str]]:
    warnings: list[str] = []
    math_stats: dict[str, int] = {}
    slide_map = {s.slide_no: s for s in slides}
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    style.font.size = BODY_SIZE

    # PW branding: extract the logo from a slide (trying several, since the title
    # slide often lacks a clean top-right mark), then apply the header (page# +
    # logo), footer, and background watermark.
    logo_path = wm_path = None
    candidates = [s.image_path for s in slides if s.image_path and Path(s.image_path).exists()]
    for img in candidates[:6]:
        logo_path, wm_path = extract_pw_logo(Path(img), run_dir / "branding")
        if logo_path:
            break
    try:
        setup_header(section, logo_path)
        setup_footer(section)
        if wm_path:
            add_watermark(section, wm_path)
    except Exception as e:
        warnings.append(f"PW branding (header/footer/watermark) could not be applied: {e}")

    add_title_block(doc, subject, chapter_title)

    inserted_images = 0
    content_added = False
    region_cache: dict = {}   # slide image -> detected diagram regions
    used_regions: dict = {}   # slide image -> how many of its regions are placed
    for raw_line in _normalize_math_tags(notes_text).splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        block = _BLOCK_ONLY_RE.match(line)
        if block:
            # Displayed equation: its own centred paragraph.
            _, used = add_block_math(
                doc, block.group(1), mode=math_render_mode, size=BODY_SIZE,
                color=DARK_GREY, fallback_writer=_math_fallback_writer,
            )
            math_stats[used] = math_stats.get(used, 0) + 1
            content_added = True
            continue
        if _is_dtp(line):
            note = parse_dtp_note(line)
            _add_dtp(doc, line, dtp_note_policy)
            if dtp_note_policy in {"keep_note_and_insert_image", "hide_note_insert_image"}:
                if note and note.slide_no and note.slide_no in slide_map and slide_map[note.slide_no].image_path:
                    img, cropped = _pick_dtp_image(
                        slide_map[note.slide_no], run_dir, region_cache, used_regions
                    )
                    if img is None:
                        warnings.append(
                            f"Slide {note.slide_no} has no unused diagram left for this note, "
                            "so the image was not repeated here."
                        )
                    else:
                        ok = _insert_slide_image(doc, img, run_dir, mode=image_insert_mode,
                                                 already_cropped=cropped)
                        if ok:
                            inserted_images += 1
                        else:
                            warnings.append(f"Could not insert image for slide {note.slide_no}.")
                else:
                    warnings.append(f"DTP note found but slide image could not be matched: {line[:140]}")
            content_added = True
            continue
        if _is_heading(line):
            if content_added:
                _add_blank_line(doc)
            _add_heading(doc, line, math_mode=math_render_mode, math_stats=math_stats)
            content_added = True
        elif _is_bullet(line):
            leading = len(raw_line) - len(raw_line.lstrip(" \t"))
            indent_level = min(3, leading // 4) if leading else 0
            _add_body(doc, line, bullet=True, indent_level=indent_level,
                      math_mode=math_render_mode, math_stats=math_stats)
            content_added = True
        else:
            _add_body(doc, line, bullet=False, math_mode=math_render_mode, math_stats=math_stats)
            content_added = True
    doc.save(str(output_path))
    warnings.append(f"Inserted {inserted_images} image(s) from DTP notes.")
    total_math = sum(math_stats.values())
    if total_math:
        parts = ", ".join(f"{n} {k}" for k, n in sorted(math_stats.items()))
        warnings.append(f"Rendered {total_math} formula(s) ({parts}).")
        if math_render_mode == RENDER_OMML and math_stats.get("unicode"):
            warnings.append(
                f"{math_stats['unicode']} formula(s) could not be built as a Word equation "
                "and fell back to Unicode text."
            )
    if not is_kalam_installed():
        warnings.append("Kalam handwritten font is not installed, so the notes render in a fallback font. Install Kalam (https://fonts.google.com/specimen/Kalam) for the handwritten PW style.")
    return output_path, warnings
