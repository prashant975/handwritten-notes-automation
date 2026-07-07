"""PW-style document furniture: title block, header, footer, and watermark.

Reproduces the Physics Wallah notes layout — subject/chapter title, PW logo in
the header, "Master NCERT with PW Books APP" footer, page numbers, and a faint
PW-logo watermark behind the text on every page.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

TITLE_COLOR = RGBColor(0xFF, 0x00, 0x00)   # PW red for subject/chapter title (#FF0000)
FOOTER_COLOR = RGBColor(0x1F, 0x4E, 0x79)  # blue-ish "Master NCERT..." link colour
BODY_FONT = "Kalam"
CHAR_SPACING_TWIPS = 30  # Expanded character spacing = 1.5 pt


def is_kalam_installed() -> bool:
    """Best-effort check for the Kalam handwritten font on Windows/Linux/Mac."""
    import glob
    import os

    dirs = [
        r"C:\Windows\Fonts",
        os.path.expandvars(r"%LOCALAPPDATA%\Microsoft\Windows\Fonts"),
        os.path.expanduser("~/.fonts"),
        os.path.expanduser("~/.local/share/fonts"),
        "/usr/share/fonts",
        "/Library/Fonts",
        os.path.expanduser("~/Library/Fonts"),
    ]
    for d in dirs:
        if d and os.path.isdir(d):
            if glob.glob(os.path.join(d, "**", "[Kk]alam*.ttf"), recursive=True):
                return True
    return False


def derive_chapter_title(stem: str) -> str:
    """Turn an uploaded file's stem into a readable chapter title.

    'Cell_-_The_Unit_of_Life_02_Class_Notes_1' -> 'Cell - The Unit of Life 02'
    """
    name = stem.replace("_", " ")
    name = re.sub(r"\s+", " ", name).strip()
    # Drop a trailing copy index like "(1)" or " 1" first.
    name = re.sub(r"\s*\(\s*\d+\s*\)\s*$", "", name).strip()
    # Then drop trailing boilerplate like "Class Notes" / "Concise Notes"
    # (and anything after it), but keep lecture numbers that are part of the title.
    name = re.sub(r"[\s_-]*\b(class|concise|handwritten|lecture)\b\s+notes\b.*$", "", name, flags=re.I).strip()
    return name.strip(" -_") or stem


def _apply_font(run, *, size, bold=False, italic=False, underline=False, color=None):
    run.font.name = BODY_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    sp = rpr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        rpr.append(sp)
    sp.set(qn("w:val"), str(CHAR_SPACING_TWIPS))


def add_title_block(doc, subject: str, chapter_title: str):
    """Subject + chapter, centered, red italic underlined (PW cover style)."""
    subj = (subject or "").strip().capitalize()
    if subj:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        _apply_font(p.add_run(subj), size=Pt(26), bold=True, italic=True, underline=True, color=TITLE_COLOR)
    if chapter_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        _apply_font(p.add_run(chapter_title), size=Pt(22), bold=True, italic=True, underline=True, color=TITLE_COLOR)


def _add_page_field(paragraph, *, bracket=True, color=None, size=Pt(11)):
    if bracket:
        _apply_font(paragraph.add_run("["), size=size, color=color)
    run = paragraph.add_run()
    _apply_font(run, size=size, color=color)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)
    if bracket:
        _apply_font(paragraph.add_run("]"), size=size, color=color)


def setup_header(section, logo_path: Path | None):
    """PW logo in the top-right corner (matches the PW notes layout)."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path and Path(logo_path).exists():
        try:
            p.add_run().add_picture(str(logo_path), height=Pt(48))
        except Exception:
            pass


def setup_footer(section):
    """'Master NCERT with PW Books APP' then a centered page number."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(p.add_run("\U0001F4D8  Master NCERT with PW Books APP"), size=Pt(12), bold=True, color=FOOTER_COLOR)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(p2, bracket=True, color=RGBColor(0x33, 0x33, 0x33))


def add_watermark(section, watermark_path: Path | None):
    """Insert a faint, centered PW logo behind the text on every page."""
    if not watermark_path or not Path(watermark_path).exists():
        return False
    try:
        from PIL import Image
        iw, ih = Image.open(watermark_path).size
    except Exception:
        iw, ih = 1, 1
    # ~1.5x the previous 0.55 page-width watermark.
    width_emu = int(section.page_width * 0.82)
    height_emu = int(width_emu * (ih / iw)) if iw else width_emu
    header = section.header
    header.is_linked_to_previous = False
    p = header.add_paragraph()
    run = p.add_run()
    try:
        run.add_picture(str(watermark_path), width=Emu(width_emu), height=Emu(height_emu))
    except Exception:
        return False
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline")) if drawing is not None else None
    if inline is None:
        return False
    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT": "0", "distB": "0", "distL": "0", "distR": "0", "simplePos": "0",
                 "relativeHeight": "0", "behindDoc": "1", "locked": "0", "layoutInCell": "1",
                 "allowOverlap": "1"}.items():
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0"); anchor.append(sp)
    posH = OxmlElement("wp:positionH"); posH.set("relativeFrom", "page")
    alignH = OxmlElement("wp:align"); alignH.text = "center"; posH.append(alignH); anchor.append(posH)
    posV = OxmlElement("wp:positionV"); posV.set("relativeFrom", "page")
    alignV = OxmlElement("wp:align"); alignV.text = "center"; posV.append(alignV); anchor.append(posV)
    if extent is not None:
        anchor.append(extent)
    eff = OxmlElement("wp:effectExtent")
    for a in ("l", "t", "r", "b"):
        eff.set(a, "0")
    anchor.append(eff)
    anchor.append(OxmlElement("wp:wrapNone"))
    if docPr is not None:
        anchor.append(docPr)
    if graphic is not None:
        anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)
    return True
